import logging
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
import pypandoc
from pathlib import Path


class GestorTablaContenido:
    """
    Clase para gestionar la inserción de tablas de contenido, índices de tablas
    e índices de figuras en documentos Word (.docx).
    """

    def __init__(self, logger: logging.Logger | None = None):
        """
        Inicializa el gestor de tablas de contenido.
        
        Args:
            logger: Logger personalizado. Si no se proporciona, se crea uno por defecto.
        """
        self.logger = logger or logging.getLogger(__name__)

    def crear_campo_toc(self, tipo: str = "contenido") -> tuple:
        """
        Crea los elementos XML para diferentes tipos de tablas de contenido.

        Args:
            tipo: Tipo de TOC a crear. Opciones: 'contenido', 'tablas', 'figuras'

        Returns:
            Tupla con los elementos XML necesarios para el campo TOC
        """
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")

        if tipo == "contenido":
            instrText.text = r'TOC \o "1-3" \h \z \u'
        elif tipo == "tablas":
            instrText.text = r'TOC \h \z \c "Tabla"'
        elif tipo == "figuras":
            instrText.text = r'TOC \h \z \c "Figura"'
        else:
            raise ValueError(f"Tipo de TOC no válido: {tipo}. Use 'contenido', 'tablas' o 'figuras'")

        fldChar_separate = OxmlElement("w:fldChar")
        fldChar_separate.set(qn("w:fldCharType"), "separate")

        updateFields = OxmlElement("w:updateFields")
        updateFields.set(qn("w:val"), "true")

        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")

        return fldChar_begin, instrText, fldChar_separate, updateFields, fldChar_end

    def insertar_toc_en_parrafo(self, paragraph, tipo: str = "contenido") -> None:
        """
        Inserta un campo TOC en un párrafo existente.

        Args:
            paragraph: Objeto paragraph de python-docx
            tipo: Tipo de TOC a insertar
        """
        run = paragraph.add_run()
        elementos = self.crear_campo_toc(tipo)

        r_element = run._r
        for elemento in elementos:
            r_element.append(elemento)

    def agregar_todas_las_toc(
        self,
        doc_path: str,
        output_path: str | None = None,
        texto_buscar: str = "PRESENTACIÓN",
        incluir_salto_pagina: bool = True
    ) -> bool:
        """
        Inserta tabla de contenido, índice de tablas e índice de figuras
        justo antes de un texto específico en el documento.

        Args:
            doc_path: Ruta del documento de entrada (.docx)
            output_path: Ruta del documento de salida. Si es None, sobreescribe el original.
            texto_buscar: Texto antes del cual insertar las TOC (default: "PRESENTACIÓN")
            incluir_salto_pagina: Si True, agrega un salto de página después de las TOC

        Returns:
            True si la operación fue exitosa, False en caso contrario
        """
        # Verificar que el archivo existe
        if not os.path.exists(doc_path):
            self.logger.error(f"No se encontró el archivo: {doc_path}")
            return False

        # Si no se especifica output, sobreescribir el original
        if output_path is None:
            output_path = doc_path

        try:
            doc = Document(doc_path)
            self.logger.info(f"Archivo cargado: {doc_path}")

            # Buscar el párrafo que contiene el texto
            posicion = None
            for i, paragraph in enumerate(doc.paragraphs):
                if texto_buscar.upper() in paragraph.text.upper():
                    posicion = i
                    self.logger.info(f"Encontrado '{texto_buscar}' en el párrafo {i}")
                    break

            # Si no se encuentra el texto, insertar al inicio
            if posicion is None:
                posicion = 0
                self.logger.warning(
                    f"No se encontró '{texto_buscar}'. Insertando al inicio del documento."
                )

            # Tabla de Contenido
            titulo_contenido = doc.paragraphs[posicion].insert_paragraph_before(
                "Tabla de Contenido"
            )
            titulo_contenido.style = "Heading 1"

            parrafo_toc = doc.paragraphs[posicion + 1].insert_paragraph_before()
            self.insertar_toc_en_parrafo(parrafo_toc, "contenido")

            doc.paragraphs[posicion + 2].insert_paragraph_before()

            # Índice de Tablas
            titulo_tablas = doc.paragraphs[posicion + 3].insert_paragraph_before(
                "Índice de Tablas"
            )
            titulo_tablas.style = "Heading 1"

            parrafo_tablas = doc.paragraphs[posicion + 4].insert_paragraph_before()
            self.insertar_toc_en_parrafo(parrafo_tablas, "tablas")

            doc.paragraphs[posicion + 5].insert_paragraph_before()

            # Índice de Figuras
            titulo_figuras = doc.paragraphs[posicion + 6].insert_paragraph_before(
                "Índice de Figuras"
            )
            titulo_figuras.style = "Heading 1"

            parrafo_figuras = doc.paragraphs[posicion + 7].insert_paragraph_before()
            self.insertar_toc_en_parrafo(parrafo_figuras, "figuras")

            doc.paragraphs[posicion + 8].insert_paragraph_before()

            # Salto de página opcional antes de PRESENTACIÓN
            if incluir_salto_pagina:
                salto_pagina = doc.paragraphs[posicion + 9].insert_paragraph_before()
                run_salto = salto_pagina.add_run()
                run_salto.add_break(WD_BREAK.PAGE)

            # Guardar documento
            doc.save(output_path)
            self.logger.info(
                f"Tablas de contenido insertadas correctamente en: {output_path}"
            )
            self.logger.info(
                "IMPORTANTE: Abre el documento en Word y presiona Ctrl+A y luego F9 para actualizar los campos"
            )
            
            return True

        except Exception as e:
            self.logger.error(f"Error al procesar el documento: {e}", exc_info=True)
            return False

    def agregar_toc_simple(
        self,
        doc_path: str,
        output_path: str | None = None,
        texto_buscar: str = "PRESENTACIÓN",
        solo_contenido: bool = False
    ) -> bool:
        """
        Versión simplificada que solo agrega la tabla de contenido principal.

        Args:
            doc_path: Ruta del documento de entrada
            output_path: Ruta de salida (None para sobreescribir)
            texto_buscar: Texto de referencia para la inserción
            solo_contenido: Si True, solo inserta tabla de contenido (sin índices)

        Returns:
            True si fue exitoso
        """
        if not os.path.exists(doc_path):
            self.logger.error(f"No se encontró el archivo: {doc_path}")
            return False

        if output_path is None:
            output_path = doc_path

        try:
            doc = Document(doc_path)
            
            # Buscar posición
            posicion = 0
            for i, paragraph in enumerate(doc.paragraphs):
                if texto_buscar.upper() in paragraph.text.upper():
                    posicion = i
                    break

            # Solo tabla de contenido
            titulo = doc.paragraphs[posicion].insert_paragraph_before("Tabla de Contenido")
            titulo.style = "Heading 1"

            parrafo_toc = doc.paragraphs[posicion + 1].insert_paragraph_before()
            self.insertar_toc_en_parrafo(parrafo_toc, "contenido")

            doc.paragraphs[posicion + 2].insert_paragraph_before()

            # Salto de página
            salto = doc.paragraphs[posicion + 3].insert_paragraph_before()
            salto.add_run().add_break(WD_BREAK.PAGE)

            doc.save(output_path)
            self.logger.info(f"Tabla de contenido insertada en: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error: {e}", exc_info=True)
            return False

    def convertir_a_pdf(self, input_docx: str, output_pdf: str | None = None) -> bool:
        """
        Convierte un archivo DOCX a PDF usando LibreOffice en modo headless.
        Funciona en Docker sin requerir pandoc ni LaTeX.
        """

        if not os.path.exists(input_docx):
            self.logger.error(f"No se encontró el archivo: {input_docx}")
            return False

        if output_pdf is None:
            output_pdf = Path(input_docx).with_suffix(".pdf")

        try:
            # LibreOffice usa como output el directorio, no el archivo final
            output_dir = str(Path(output_pdf).parent)

            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", output_dir,
                    input_docx
                ],
                check=True
            )

            self.logger.info(f"Archivo convertido a PDF: {output_pdf}")
            return True

        except Exception as e:
            self.logger.error(f"Error al convertir a PDF: {e}", exc_info=True)
            return False


