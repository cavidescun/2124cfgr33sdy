import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
import pypandoc
from docxcompose.composer import Composer


def crear_campo_toc(tipo="contenido"):
    """Crea los elementos XML para diferentes tipos de tablas de contenido."""
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")

    if tipo == "contenido":
        instrText.text = r'TOC \o "1-3" \h \z \u'
    elif tipo == "tablas":
        instrText.text = r'TOC \h \z \c "Tabla"'
    elif tipo == "figuras":
        instrText.text = r'TOC \h \z \c "Figura"'

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    updateFields = OxmlElement("w:updateFields")
    updateFields.set(qn("w:val"), "true")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    return fldChar_begin, instrText, fldChar_separate, updateFields, fldChar_end


def insertar_toc_en_parrafo(paragraph, tipo="contenido"):
    """Inserta un campo TOC en un párrafo existente."""
    run = paragraph.add_run()
    for elemento in crear_campo_toc(tipo):
        run._r.append(elemento)


def unir_documentos_por_llave(base_path, llave_maestra, archivo_output):
    carpeta_llave = os.path.join(base_path, llave_maestra)

    if not os.path.exists(carpeta_llave):
        print(f"❌ No se encontró la carpeta para la llave '{llave_maestra}'")
        return None

    archivos_docx = [f for f in os.listdir(carpeta_llave) if f.endswith(".docx")]
    if not archivos_docx:
        print(f"❌ No se encontraron archivos .docx en {carpeta_llave}")
        return None

    def clave_orden(nombre):
        base = os.path.splitext(nombre)[0]
        return (not base.isdigit(), int(base) if base.isdigit() else base.lower())

    archivos_docx = sorted(archivos_docx, key=clave_orden)

    print(f"📄 Archivos encontrados ({len(archivos_docx)}):")
    for i, archivo in enumerate(archivos_docx, 1):
        print(f"   {i}. {archivo}")

    doc_unificado = Document(os.path.join(carpeta_llave, archivos_docx[0]))
    composer = Composer(doc_unificado)

    for archivo in archivos_docx[1:]:
        ruta = os.path.join(carpeta_llave, archivo)
        doc_temp = Document(ruta)
        composer.append(doc_temp)
        print(f"   ✓ Añadido: {archivo}")

    composer.save(archivo_output)
    print(f"\n✅ Documentos unidos correctamente en: {archivo_output}")
    return archivo_output


def agregar_todas_las_toc(doc_path, output_path, texto_buscar="PRESENTACIÓN"):
    """Inserta tabla de contenido, índice de tablas y figuras antes del texto indicado."""
    doc = Document(doc_path)
    posicion = None

    for i, p in enumerate(doc.paragraphs):
        if texto_buscar.upper() in p.text.upper():
            posicion = i
            print(f"✓ Encontrado '{texto_buscar}' en párrafo {i}")
            break

    if posicion is None:
        posicion = 0
        print(f"⚠️ No se encontró '{texto_buscar}'. Se insertará al inicio.")

    titulo_contenido = doc.paragraphs[posicion].insert_paragraph_before(
        "Tabla de Contenido"
    )
    titulo_contenido.style = "Heading 1"
    parrafo_toc = doc.paragraphs[posicion + 1].insert_paragraph_before()
    insertar_toc_en_parrafo(parrafo_toc, "contenido")

    titulo_tablas = doc.paragraphs[posicion + 2].insert_paragraph_before(
        "Índice de Tablas"
    )
    titulo_tablas.style = "Heading 1"
    parrafo_tablas = doc.paragraphs[posicion + 3].insert_paragraph_before()
    insertar_toc_en_parrafo(parrafo_tablas, "tablas")

    titulo_figuras = doc.paragraphs[posicion + 4].insert_paragraph_before(
        "Índice de Figuras"
    )
    titulo_figuras.style = "Heading 1"
    parrafo_figuras = doc.paragraphs[posicion + 5].insert_paragraph_before()
    insertar_toc_en_parrafo(parrafo_figuras, "figuras")

    salto = doc.paragraphs[posicion + 6].insert_paragraph_before()
    run_salto = salto.add_run()
    run_salto.add_break(WD_BREAK.PAGE)

    doc.save(output_path)
    print(f"✅ TOCs insertadas antes de '{texto_buscar}' en {output_path}")


def convertir_a_pdf(input_docx, output_pdf):
    """Convierte un archivo DOCX a PDF."""
    try:
        pypandoc.convert_file(
            input_docx, "pdf", outputfile=output_pdf, extra_args=["--standalone"]
        )
        print(f"✅ Archivo convertido a PDF: {output_pdf}")
    except Exception as e:
        print(f"⚠️ Error al convertir a PDF: {e}")


if __name__ == "__main__":
    base_path = "output"  # Carpeta raíz donde están las llaves
    llave_maestra = input("🔑 Ingresa la llave maestra: ").strip()

    archivo_unificado = f"{llave_maestra}_unificado.docx"
    archivo_final = f"{llave_maestra}_final_con_toc.docx"
    archivo_pdf = f"{llave_maestra}_final.pdf"

    TEXTO_REFERENCIA = "PRESENTACIÓN"

    print("=" * 60)
    print("UNIR DOCUMENTOS WORD POR LLAVE MAESTRA Y CREAR TOC")
    print("=" * 60)

    resultado = unir_documentos_por_llave(base_path, llave_maestra, archivo_unificado)

    if resultado:
        print("\n📑 Insertando tablas de contenido...")
        agregar_todas_las_toc(archivo_unificado, archivo_final, TEXTO_REFERENCIA)

        resp = input("\n¿Deseas convertir a PDF? (s/n): ").lower()
        if resp == "s":
            convertir_a_pdf(archivo_final, archivo_pdf)

        print("\n✅ Proceso completo.")
