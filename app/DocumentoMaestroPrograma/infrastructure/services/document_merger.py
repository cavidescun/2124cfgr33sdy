import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer


class DocumentMerger:
    """Clase para unir documentos Word y agregar tablas de contenido."""

    def __init__(self, base_path="output", verbose=True):
        """
        Inicializa el DocumentMerger.

        Args:
            base_path (str): Carpeta raíz donde están las llaves
            verbose (bool): Si True, imprime mensajes de progreso
        """
        self.base_path = base_path
        self.verbose = verbose

    def _print(self, mensaje):
        """Imprime mensaje solo si verbose está activado."""
        if self.verbose:
            print(mensaje)

    def _crear_campo_toc(self, tipo="contenido"):
        """
        Crea los elementos XML para diferentes tipos de tablas de contenido.

        Args:
            tipo (str): Tipo de TOC ('contenido', 'tablas', 'figuras')

        Returns:
            tuple: Elementos XML para el campo TOC
        """
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")

        if tipo == "contenido":
            instrText.text = r'TOC \o "1-3" \h \z \u'
        elif tipo == "tablas":
            instrText.text = r'TOC \h \z \c "Tabla"'
        elif tipo == "figuras":
            instrText.text = r'TOC \h \z \c "Figura"'

        fldChar_separate = OxmlElement("w:fldChar")
        fldChar_separate.set(qn("w:fldCharType"), "separate")

        updateFields = OxmlElement("w:updateFields")
        updateFields.set(qn("w:val"), "true")

        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")

        return fldChar_begin, instrText, fldChar_separate, updateFields, fldChar_end

    def _insertar_toc_en_parrafo(self, paragraph, tipo="contenido"):
        """
        Inserta un campo TOC en un párrafo existente.

        Args:
            paragraph: Párrafo de Word donde insertar el TOC
            tipo (str): Tipo de TOC a insertar
        """
        run = paragraph.add_run()
        for elemento in self._crear_campo_toc(tipo):
            run._r.append(elemento)

    def unir_documentos(self, llave_maestra, archivo_output):
        """
        Une todos los documentos .docx de una carpeta específica.

        Args:
            llave_maestra (str): Nombre de la carpeta/llave que contiene los documentos
            archivo_output (str): Ruta del archivo de salida

        Returns:
            str: Ruta del archivo generado o None si hay error
        """
        carpeta_llave = os.path.join(self.base_path, llave_maestra)

        if not os.path.exists(carpeta_llave):
            self._print(f"❌ No se encontró la carpeta para la llave '{llave_maestra}'")
            return None

        archivos_docx = [f for f in os.listdir(carpeta_llave) if f.endswith(".docx")]
        if not archivos_docx:
            self._print(f"❌ No se encontraron archivos .docx en {carpeta_llave}")
            return None

        def clave_orden(nombre):
            base = os.path.splitext(nombre)[0]
            return (not base.isdigit(), int(base) if base.isdigit() else base.lower())

        archivos_docx = sorted(archivos_docx, key=clave_orden)

        self._print(f"📄 Archivos encontrados ({len(archivos_docx)}):")
        for i, archivo in enumerate(archivos_docx, 1):
            self._print(f"   {i}. {archivo}")

        doc_unificado = Document(os.path.join(carpeta_llave, archivos_docx[0]))
        composer = Composer(doc_unificado)

        for archivo in archivos_docx[1:]:
            ruta = os.path.join(carpeta_llave, archivo)
            doc_temp = Document(ruta)
            composer.append(doc_temp)
            self._print(f"   ✓ Añadido: {archivo}")

        composer.save(archivo_output)
        self._print(f"\n✅ Documentos unidos correctamente en: {archivo_output}")
        return archivo_output

    def agregar_toc(self, doc_path, output_path, texto_buscar="PRESENTACIÓN"):
        """
        Inserta tabla de contenido, índice de tablas y figuras antes del texto indicado.

        Args:
            doc_path (str): Ruta del documento de entrada
            output_path (str): Ruta del documento de salida
            texto_buscar (str): Texto de referencia para insertar las TOCs

        Returns:
            str: Ruta del archivo generado
        """
        doc = Document(doc_path)
        posicion = None

        for i, p in enumerate(doc.paragraphs):
            if texto_buscar.upper() in p.text.upper():
                posicion = i
                self._print(f"✓ Encontrado '{texto_buscar}' en párrafo {i}")
                break

        if posicion is None:
            posicion = 0
            self._print(f"⚠️ No se encontró '{texto_buscar}'. Se insertará al inicio.")

        titulo_contenido = doc.paragraphs[posicion].insert_paragraph_before(
            "Tabla de Contenido"
        )
        titulo_contenido.style = "Heading 1"
        parrafo_toc = doc.paragraphs[posicion + 1].insert_paragraph_before()
        self._insertar_toc_en_parrafo(parrafo_toc, "contenido")

        titulo_tablas = doc.paragraphs[posicion + 2].insert_paragraph_before(
            "Índice de Tablas"
        )
        titulo_tablas.style = "Heading 1"
        parrafo_tablas = doc.paragraphs[posicion + 3].insert_paragraph_before()
        self._insertar_toc_en_parrafo(parrafo_tablas, "tablas")

        titulo_figuras = doc.paragraphs[posicion + 4].insert_paragraph_before(
            "Índice de Figuras"
        )
        titulo_figuras.style = "Heading 1"
        parrafo_figuras = doc.paragraphs[posicion + 5].insert_paragraph_before()
        self._insertar_toc_en_parrafo(parrafo_figuras, "figuras")

        salto = doc.paragraphs[posicion + 6].insert_paragraph_before()
        run_salto = salto.add_run()
        run_salto.add_break(WD_BREAK.PAGE)

        doc.save(output_path)
        self._print(f"✅ TOCs insertadas antes de '{texto_buscar}' en {output_path}")
        return output_path

    def procesar_completo(self, llave_maestra, texto_referencia="PRESENTACIÓN"):
        """
        Ejecuta el proceso completo: unir documentos y agregar TOCs.
        El archivo final se guarda en output/LLAVE/reporte/

        Args:
            llave_maestra (str): Nombre de la carpeta/llave
            texto_referencia (str): Texto de referencia para las TOCs

        Returns:
            dict: Diccionario con las rutas de los archivos generados
        """
        carpeta_reporte = os.path.join(self.base_path, llave_maestra, "reporte")
        os.makedirs(carpeta_reporte, exist_ok=True)
        archivo_unificado_temp = f"{llave_maestra}_unificado_temp.docx"
        
        archivo_final = os.path.join(carpeta_reporte, f"{llave_maestra}_final_con_toc.docx")

        self._print("=" * 60)
        self._print("UNIR DOCUMENTOS WORD POR LLAVE MAESTRA Y CREAR TOC")
        self._print("=" * 60)
        self._print(f"📁 Carpeta de salida: {carpeta_reporte}")

        resultado = self.unir_documentos(llave_maestra, archivo_unificado_temp)

        if resultado:
            self._print("\n📑 Insertando tablas de contenido...")
            self.agregar_toc(archivo_unificado_temp, archivo_final, texto_referencia)
            
            if os.path.exists(archivo_unificado_temp):
                os.remove(archivo_unificado_temp)
                self._print(f"🗑️  Archivo temporal eliminado: {archivo_unificado_temp}")
            
            self._print("\n✅ Proceso completo.")

            return {
                "final": archivo_final,
                "carpeta_reporte": carpeta_reporte,
                "success": True,
            }

        return {"success": False}